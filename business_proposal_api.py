from fastapi import FastAPI, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.responses import StreamingResponse
from starlette.background import BackgroundTask

from docx import Document
from docx.shared import Inches
import os
import requests

import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import os
import tempfile

app = FastAPI()

# Allow CORS for frontend communication
origins = [
    "http://localhost",
    "http://localhost:3000",  # Frontend development environment
    "*"  # Allow all origins for testing purposes; restrict in production
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define required topics
required_topics = [
    "Company Overview", "Mission and Vision Statement", "Executive Summary",
    "Owners and Partnerships", "Industry Overview and Trends", "Competition",
    "Problem Statement", "Marketing Plan", "Proposed Solution", "Market Analysis",
    "Sustainable Practices", "Implementation Timeline", "Staff Names",
    "Financial Objectives", "Exit Strategy", "Conclusion"
]

# Define proposal templates
templates = {
    "Template 1": [
        "Company Overview", "Mission and Vision Statement", "Executive Summary",
        "Problem Statement", "Proposed Solution", "Market Analysis",
        "Marketing Plan", "Implementation Timeline", "Conclusion"
    ],
    "Template 2": [
        "Executive Summary", "Problem Statement", "Market Analysis",
        "Proposed Solution", "Marketing Plan", "Implementation Timeline",
        "Conclusion", "Company Overview", "Mission and Vision Statement"
    ],
    "Template 3": [
        "Market Analysis", "Problem Statement", "Proposed Solution",
        "Marketing Plan", "Implementation Timeline", "Conclusion",
        "Mission and Vision Statement", "Company Overview", "Executive Summary"
    ]
}

# OCR.space API details
OCR_SPACE_API_URL = "https://api.ocr.space/parse/image"
OCR_SPACE_API_KEY = "K84204290188957"  # Replace with your OCR.space API key

# Groq API details
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_API_KEY = "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"  # Replace with your Groq API key

def extract_text_from_pdf(file: UploadFile):
    """
    Extracts text from the uploaded PDF using OCR.space and translates Sinhala text to English if detected.
    Returns the extracted text, detected language, and translation status.
    """
    # Step 1: Call OCR.space API to extract text
    response = requests.post(
        OCR_SPACE_API_URL,
        files={'file': (file.filename, file.file)},
        data={'apikey': OCR_SPACE_API_KEY, 'language': 'eng'}
    )

    # Check for a successful response
    if response.status_code != 200:
        raise Exception(f"OCR API Error: {response.status_code} - {response.text}")

    try:
        result = response.json()
        extracted_text = result.get('ParsedResults', [{}])[0].get('ParsedText', '')
        print("Extracted Text:", extracted_text)  # Debug print
    except ValueError as e:
        raise ValueError(f"Failed to parse JSON response: {response.text}") from e

    # Step 2: Detect the language of the extracted text
    from googletrans import Translator
    translator = Translator()

    try:
        detected_lang = translator.detect(extracted_text).lang
        print("Detected Language:", detected_lang)  # Debug print
    except Exception as e:
        raise Exception(f"Language detection failed: {e}")

    # Step 3: Override language detection if Sinhala text is found
    if detected_lang == 'en':  # Detected as English
        # Check for Sinhala Unicode characters in the text
        if any('\u0D80' <= char <= '\u0DFF' for char in extracted_text):  # Sinhala Unicode range
            print("Overriding detected language to Sinhala (si)")  # Debug print
            detected_lang = 'si'

    # Step 4: Translate text to English if detected language is Sinhala
    translation_status = False
    if detected_lang == 'si':
        print("Detected Sinhala text. Translating to English...")
        try:
            extracted_text = translator.translate(extracted_text, src='si', dest='en').text
            translation_status = True
        except Exception as e:
            raise Exception(f"Translation failed: {e}")
        print("Translation completed.")

    return {
        "text": extracted_text,
        "language": detected_lang,
        "translated": translation_status
    }

def identify_missing_topics(extracted_text):
    missing = []
    for topic in required_topics:
        if topic.lower() not in extracted_text.lower():
            missing.append(topic)
    return missing

def generate_proposal_with_groq(metadata, user_inputs, topic_order):
    content = {}
    for topic in topic_order:
        details = user_inputs.get(topic, "")
        prompt = (
            f"Create detailed content for the topic '{topic}'.\n"
            f"Business Name: {metadata['name']}\n"
            f"Domain: {metadata['domain']}\n"
            f"Existing Business: {metadata['is_existing']}\n"
            f"Instructions: {metadata['user_instructions']}\n"
            f"Details: {details}\n"
        )

        response = requests.post(
            GROQ_API_URL,
            headers={'Authorization': f'Bearer {GROQ_API_KEY}'},
            json={
                "messages": [{"role": "user", "content": prompt}],
                "model": "llama-3.1-8b-instant"
            }
        )
        result = response.json()
        topic_content = result['choices'][0]['message']['content']
        content[topic] = topic_content

    return content

def save_proposal_to_word(content, output_path):
    doc = Document()
    doc.add_heading("Business Proposal", level=0)
    for topic, details in content.items():
        doc.add_heading(topic, level=1)
        doc.add_paragraph(details)
        doc.add_page_break()
    doc.save(output_path)

def sinhala_to_english_translation(input_text):
    translator = Translator()
    translation = translator.translate(input_text, src='si', dest='en')
    return translation.text

@app.post("/check-missing-topics/")
async def check_missing_topics(file: UploadFile = File(...)):
    # Extract text and get metadata
    extraction_result = extract_text_from_pdf(file)
    extracted_text = extraction_result["text"]
    detected_lang = extraction_result["language"]
    translation_status = extraction_result["translated"]

    # Identify missing topics
    missing_topics = identify_missing_topics(extracted_text)

    return {
        "missing_info": "yes" if missing_topics else "no",
        "missing_topics": missing_topics,
        "detected_language": detected_lang,
        "translated_to_english": translation_status
    }

def extract_headings_and_details(doc_path):
    from docx import Document

    doc = Document(doc_path)
    content = []
    current_heading = None

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading 1"):
            if current_heading:
                content.append(current_heading)
            current_heading = {"heading": paragraph.text, "details": ""}
        elif current_heading:
            current_heading["details"] += paragraph.text + "\n"

    if current_heading:
        content.append(current_heading)

    return content

@app.post("/generate-proposal/")
async def generate_proposal(
    file: UploadFile = File(...),
    business_name: str = Form(...),
    business_domain: str = Form(...),
    is_existing: str = Form(...),
    user_instructions: str = Form(...),
    selected_template: str = Form(...),
):
    # Step 1: Extract text from PDF
    extraction_result = extract_text_from_pdf(file)
    extracted_text = extraction_result["text"]  # Get only the text part
    detected_lang = extraction_result["language"]
    translation_status = extraction_result["translated"]

    # Step 2: Identify missing topics
    missing_topics = identify_missing_topics(extracted_text)

    # Step 3: Collect user inputs for missing topics
    user_inputs = {topic: "Details not provided" for topic in missing_topics}

    # Step 4: Metadata for proposal generation
    metadata = {
        "name": business_name,
        "domain": business_domain,
        "is_existing": is_existing,
        "user_instructions": user_instructions,
    }

    # Step 5: Generate proposal using the selected template
    topic_order = templates.get(selected_template, required_topics)
    proposal_content = generate_proposal_with_groq(metadata, user_inputs, topic_order)

    # Step 6: Save the proposal to a Word document
    output_path = os.path.join(os.getcwd(), "Business_Proposal.docx")
    save_proposal_to_word(proposal_content, output_path)

    # Step 7: Prepare the document for streaming response
    def cleanup_file():
        try:
            os.remove(output_path)
        except OSError as e:
            print(f"Error deleting file: {e}")

    file_response = StreamingResponse(
        open(output_path, "rb"),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Business_Proposal.docx"},
        background=BackgroundTask(cleanup_file),
    )

    # Step 8: Return JSON with file response and translation info
    return JSONResponse(content={
        "headings": extract_headings_and_details(output_path),
        "detected_language": detected_lang,
        "translated_to_english": translation_status,
        "download_url": "/download-proposal/"
    })


@app.post("/download-proposal/")
async def download_proposal(image: UploadFile = File(...)):
    """
    Accepts an image from the frontend, adds it as the first page (cover page)
    in the Business_Proposal.docx, and sends the modified document to the frontend.
    """

    # Define paths
    proposal_path = os.path.join(os.getcwd(), "Business_Proposal.docx")
    modified_proposal_path = os.path.join(os.getcwd(), "Modified_Business_Proposal.docx")

    # Ensure the original proposal exists
    if not os.path.exists(proposal_path):
        return {"error": "Business_Proposal.docx not found"}

    # Open the existing proposal document
    doc = Document()

    # Step 1: Save the received image
    image_path = os.path.join(os.getcwd(), "cover_page.png")
    with open(image_path, "wb") as buffer:
        buffer.write(await image.read())

    # Step 2: Insert image as the first page (cover page)
    doc.add_picture(image_path, width=Inches(6))  # Adjust width as needed
    doc.add_page_break()  # Add a page break after the cover page

    # Step 3: Append content from the existing proposal
    existing_doc = Document(proposal_path)
    for element in existing_doc.element.body:
        doc.element.body.append(element)

    # Step 4: Save the modified document
    doc.save(modified_proposal_path)

    # Step 5: Cleanup function to delete temporary files after response
    def cleanup_files():
        try:
            os.remove(image_path)
            os.remove(modified_proposal_path)
        except OSError as e:
            print(f"Error deleting file: {e}")

    # Step 6: Send the modified document to the frontend
    return FileResponse(
        modified_proposal_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Modified_Business_Proposal.docx",
        background=BackgroundTask(cleanup_files)
    )