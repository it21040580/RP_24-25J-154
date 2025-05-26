from fastapi import FastAPI, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from docx import Document
import os
import requests

app = FastAPI()

# Allow CORS for frontend communication
origins = [
    "http://localhost",
    "http://localhost:3000",  # Frontend development environment
    "*"  # Allow all origins for testing purposes; restrict in production
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define required topics
required_topics = [
    "Company Overview", "Mission and Vision Statement", "Executive Summary",
    "Owners and Partnerships", "Industry Overview and Trends", "Competition",
    "Problem Statement", "Marketing Plan", "Proposed Solution", "Market Analysis",
    "Sustainable Practices", "Implementation Timeline", "Staff Names",
    "Financial Objectives", "Exit Strategy", "Conclusion"
]

# Define proposal templates
templates = {
    "Template 1": [
        "Company Overview", "Mission and Vision Statement", "Executive Summary",
        "Problem Statement", "Proposed Solution", "Market Analysis",
        "Marketing Plan", "Implementation Timeline", "Conclusion"
    ],
    "Template 2": [
        "Executive Summary", "Problem Statement", "Market Analysis",
        "Proposed Solution", "Marketing Plan", "Implementation Timeline",
        "Conclusion", "Company Overview", "Mission and Vision Statement"
    ],
    "Template 3": [
        "Market Analysis", "Problem Statement", "Proposed Solution",
        "Marketing Plan", "Implementation Timeline", "Conclusion",
        "Mission and Vision Statement", "Company Overview", "Executive Summary"
    ]
}

# OCR.space API details
OCR_SPACE_API_URL = "https://api.ocr.space/parse/image"
OCR_SPACE_API_KEY = "K84886360488957"  # Replace with your OCR.space API key

# Groq API details
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_API_KEY = "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"  # Replace with your Groq API key

def extract_text_from_pdf(file: UploadFile):
    response = requests.post(
        OCR_SPACE_API_URL,
        files={'file': (file.filename, file.file)},
        data={'apikey': OCR_SPACE_API_KEY, 'language': 'eng'}
    )
    
    # Check for a successful response
    if response.status_code != 200:
        raise Exception(f"OCR API Error: {response.status_code} - {response.text}")
    
    try:
        result = response.json()
        return result.get('ParsedResults', [{}])[0].get('ParsedText', '')
    except ValueError as e:
        raise ValueError(f"Failed to parse JSON response: {response.text}") from e

def identify_missing_topics(extracted_text):
    missing = []
    for topic in required_topics:
        if topic.lower() not in extracted_text.lower():
            missing.append(topic)
    return missing

def generate_proposal_with_groq(metadata, user_inputs, topic_order):
    content = {}
    for topic in topic_order:
        details = user_inputs.get(topic, "")
        prompt = (
            f"Create detailed content for the topic '{topic}'.\n"
            f"Business Name: {metadata['name']}\n"
            f"Domain: {metadata['domain']}\n"
            f"Existing Business: {metadata['is_existing']}\n"
            f"Instructions: {metadata['user_instructions']}\n"
            f"Details: {details}\n"
        )

        response = requests.post(
            GROQ_API_URL,
            headers={'Authorization': f'Bearer {GROQ_API_KEY}'},
            json={
                "messages": [{"role": "user", "content": prompt}],
                "model": "llama-3.1-8b-instant"
            }
        )
        result = response.json()
        topic_content = result['choices'][0]['message']['content']
        content[topic] = topic_content

    return content

def save_proposal_to_word(content, output_path):
    doc = Document()
    doc.add_heading("Business Proposal", level=0)
    for topic, details in content.items():
        doc.add_heading(topic, level=1)
        doc.add_paragraph(details)
        doc.add_page_break()
    doc.save(output_path)

@app.post("/check-missing-topics/")
async def check_missing_topics(file: UploadFile = File(...)):
    # Extract text from PDF
    extracted_text = extract_text_from_pdf(file)

    # Identify missing topics
    missing_topics = identify_missing_topics(extracted_text)

    if missing_topics:
        return {"missing_info": "yes", "missing_topics": missing_topics}
    else:
        return {"missing_info": "no", "missing_topics": []}

@app.post("/generate-proposal/")
async def generate_proposal(
    file: UploadFile = File(...),
    business_name: str = Form(...),
    business_domain: str = Form(...),
    is_existing: str = Form(...),
    user_instructions: str = Form(...),
    selected_template: str = Form(...)
):
    # Step 1: Extract text from PDF
    extracted_text = extract_text_from_pdf(file)

    # Step 2: Identify missing topics
    missing_topics = identify_missing_topics(extracted_text)

    # Step 3: Collect user inputs for missing topics or set dummy values
    user_inputs = {topic: "Details not provided" for topic in missing_topics}

    # Step 4: Collect business metadata
    metadata = {
        "name": business_name,
        "domain": business_domain,
        "is_existing": is_existing,
        "user_instructions": user_instructions,
    }

    # Step 5: Determine topic order based on selected template
    topic_order = templates.get(selected_template, required_topics)

    # Step 6: Generate business proposal using Groq API
    proposal_content = generate_proposal_with_groq(metadata, user_inputs, topic_order)

    # Step 7: Save proposal to Word document
    output_path = os.path.join(os.getcwd(), "Business_Proposal.docx")
    save_proposal_to_word(proposal_content, output_path)

    # Step 8: Return the generated document directly as a downloadable file
    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Business_Proposal.docx"
    )